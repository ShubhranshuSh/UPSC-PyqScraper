import io
import re
from datetime import datetime
from pathlib import Path

import streamlit as st
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

from scraper import SUBJECTS, scrape


st.set_page_config(page_title="UPSC PYQ Scraper", layout="centered")

INVALID_XML_RE = re.compile(r"[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]")


def clean_text(value) -> str:
    if value is None:
        return ""
    return INVALID_XML_RE.sub("", str(value))


def build_word_doc(questions: list, year: int, subject: str) -> bytes:
    doc = Document()
    title = doc.add_heading(f"UPSC Prelims {year} - {clean_text(subject)}", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph(
        f"Total Questions: {len(questions)} | Generated: {datetime.now().strftime('%d %b %Y %H:%M')}"
    )
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")

    for q in questions:
        difficulty = clean_text(q.get("difficulty", ""))
        question = clean_text(q.get("question", ""))
        heading = f"Q{q.get('number', 0)}"
        if difficulty:
            heading += f" [{difficulty}]"
        doc.add_heading(heading, level=2)
        doc.add_paragraph(question)

        for key in ["A", "B", "C", "D"]:
            if key in q.get("options", {}):
                option_line = f"{key}. {clean_text(q['options'][key])}"
                p = doc.add_paragraph(option_line)
                p.paragraph_format.left_indent = Inches(0.3)
                if key == q.get("answer") and p.runs:
                    p.runs[0].bold = True

        if q.get("answer"):
            doc.add_paragraph(f"Correct Answer: Option {clean_text(q['answer'])}")

        explanation = clean_text(q.get("explanation", ""))
        if explanation:
            doc.add_paragraph("Explanation:")
            doc.add_paragraph(explanation)

        doc.add_paragraph("-" * 50)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


def save_doc_to_subject_folder(word_bytes: bytes, subject: str, year: int) -> Path:
    subject_slug = subject.lower().replace("&", "and").replace(" ", "_")
    subject_folder = Path("output") / subject_slug
    subject_folder.mkdir(parents=True, exist_ok=True)

    file_path = subject_folder / f"{year}.docx"
    file_path.write_bytes(word_bytes)
    return file_path.resolve()


st.title("UPSC PYQ Scraper")
st.write("Scrape previous year questions and export them to Word.")

col1, col2 = st.columns([2, 1])
with col1:
    selected_subject = st.selectbox("Subject", options=list(SUBJECTS.keys()), index=0)
with col2:
    selected_year = st.number_input("Year", min_value=2011, max_value=2025, value=2023, step=1)

if st.button("Start Scraping", use_container_width=True):
    try:
        with st.spinner("Fetching questions..."):
            questions, _ = scrape(int(selected_year), selected_subject)
    except Exception as exc:
        st.error(f"Error: {exc}")
        questions = []

    if not questions:
        st.warning("No questions found for this subject/year.")
    else:
        st.success(f"Fetched {len(questions)} questions for {selected_subject} {selected_year}.")

        word_bytes = build_word_doc(questions, int(selected_year), selected_subject)
        save_doc_to_subject_folder(word_bytes, selected_subject, int(selected_year))

        safe_subject_name = selected_subject.replace(" ", "_").replace("&", "and")
        filename = f"UPSC_PYQ_{safe_subject_name}_{selected_year}.docx"
        st.download_button(
            label="Download PYQ Doc",
            data=word_bytes,
            file_name=filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )

        st.subheader("Preview")
        for q in questions:
            st.write(f"Q{q.get('number', 0)} [{q.get('difficulty', '')}]")
            st.write(clean_text(q.get("question", "")))

            for key in ["A", "B", "C", "D"]:
                if key in q.get("options", {}):
                    text = clean_text(q["options"][key])
                    if key == q.get("answer"):
                        st.write(f"**{key}. {text}**")
                    else:
                        st.write(f"{key}. {text}")

            if q.get("explanation"):
                with st.expander("Explanation"):
                    st.write(clean_text(q["explanation"]))

            st.write("---")